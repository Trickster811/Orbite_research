from docx import Document
from docx.shared import Pt

# Création du document Word
doc = Document()

# Style du document
# style = doc.styles['Normal']
# font = style.font
# font.name = 'Calibri'
# font.size = Pt(11)

# Titre principal
doc.add_heading('Correction de l’Examen – Système d’Exploitation', 0)

# Exercice 1
doc.add_heading('Exercice 1 : Questions de Compréhension', level=1)
questions1 = [
    ("Je suis un appel système qui affiche le PID du processus père", "getppid()"),
    ("Je suis un appel système qui permet de créer un processus", "fork()"),
    ("Je suis un processus qui s'est terminé mais dont le père n'a pas encore lu le code de retour", "Processus zombie"),
    ("Je suis la différence entre le temps de la première exécution et le temps d'entrée dans le système", "Temps de réponse"),
    ("Je suis la différence entre le temps de terminaison et le temps d'entrée dans le système", "Temps de séjour"),
    ("Je suis le premier programme lancé à la mise sous tension de l’ordinateur", "Le BIOS ou le chargeur de démarrage (bootloader)"),
]
for q, a in questions1:
    doc.add_paragraph(f"{q} : ", style='List Bullet').add_run(f" {a}").bold = True

# Exercice 2
doc.add_heading('Exercice 2 : Questions de Cours', level=1)

doc.add_heading('1. Différence entre un processeur et un processus', level=2)
doc.add_paragraph("- Processeur (CPU) : composant matériel qui exécute les instructions machine.")
doc.add_paragraph("- Processus : instance logicielle d’un programme en cours d'exécution, avec son propre contexte, mémoire, ressources, etc.")

doc.add_heading('2. Définition d’un ordonnanceur', level=2)
doc.add_paragraph("Un ordonnanceur (scheduler) est une partie du système d’exploitation chargée de décider quel processus sera exécuté à un instant donné, selon une politique d’ordonnancement :")
doc.add_paragraph("• FIFO (First In First Out)\n• Priorité\n• Round-Robin\n• SJF (Shortest Job First)", style='List Bullet')

doc.add_heading('3. Microprogramme de gestion des informations système', level=2)
doc.add_paragraph("a. Affichage des informations système et machine :")
doc.add_paragraph("uname -a\nou\nlshw", style='Intense Quote')
doc.add_paragraph("b. Enregistrement de ces informations dans un fichier texte :")
doc.add_paragraph("uname -a > infos.txt", style='Intense Quote')
doc.add_paragraph("c. Ajout de la date actuelle à ce fichier :")
doc.add_paragraph("date >> infos.txt", style='Intense Quote')

# Exercice 3
doc.add_heading('Exercice 3 : Ordonnancement', level=1)
doc.add_paragraph("Voici une structure type à suivre pour répondre à ce type d’exercice :")

doc.add_heading('a. Données du problème', level=2)
doc.add_paragraph("Liste des processus avec leur temps d’arrivée et durée d’exécution (burst time).")
doc.add_paragraph("| Processus | Temps d’arrivée | Durée d’exécution |\n|----------|-----------------|-------------------|\n| P1       | ...             | ...               |\n| P2       | ...             | ...               |")

doc.add_heading('b. Diagramme de Gantt', level=2)
doc.add_paragraph("| 0 | P1 |     | P2 |     | ...")

doc.add_heading('c. Calculs des temps', level=2)
doc.add_paragraph("- Temps de réponse = Début d'exécution - Temps d’arrivée")
doc.add_paragraph("- Temps de séjour = Fin d’exécution - Temps d’arrivée")
doc.add_paragraph("- Temps d’attente = Temps de séjour - Durée d’exécution")
doc.add_paragraph("| Processus | Réponse | Séjour | Attente |\n|----------|---------|--------|---------|\n| P1       | ...     | ...    | ...     |\n| P2       | ...     | ...    | ...     |")

doc.add_heading('d. Moyennes', level=2)
doc.add_paragraph("- Moyenne des temps de réponse")
doc.add_paragraph("- Moyenne des temps de séjour")
doc.add_paragraph("- Moyenne des temps d’attente")

# Sauvegarde du document
output_path = "Correction_Complete_Examen_SE.docx"
doc.save(output_path)

